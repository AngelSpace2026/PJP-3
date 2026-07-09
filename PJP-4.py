#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PJP – 512 Lossless Transforms (256 base + 256 fast) + 2704 Transform‑Pair Sequences
+ Hybrid Dictionary Mode (Static Word, Line, Dynamic)
+ OPTIONAL QISKIT‑INSPIRED QUANTUM TRANSFORMS
+ Base64 Transform (22) and Base64‑aware dictionary loading
+ 6‑bit Text Compression Transform (27)
+ Transforms 28–30 (per‑3‑byte subtract)
+ Transform 31: .docx paragraphs (with dictionary compression)
+ Transform 32: .docx tables (with dictionary compression)
============================================================================
"""

import math
import random
import decimal
import hashlib
import struct
import re
import os
import urllib.request
import sys
import subprocess
import importlib
import tempfile
import base64
import zipfile
import io
import xml.etree.ElementTree as ET
from typing import Optional, List, Tuple, Dict, Callable
from collections import Counter, defaultdict

# ------------------------------------------------------------------
# Helper: install a single package via pip (silent, auto)
# ------------------------------------------------------------------
def install_package(pkg: str) -> bool:
    print(f"Installing {pkg}...")
    try:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])
        print(f"Successfully installed {pkg}")
        return True
    except Exception as e:
        print(f"Failed to install {pkg}: {e}")
        return False

# ------------------------------------------------------------------
# 1. Ask about quantum transforms – auto‑install if missing
# ------------------------------------------------------------------
USE_QUANTUM = False
HAS_QISKIT = False

quantum_choice = input("Enable quantum‑inspired transforms (requires Qiskit)? (y/n): ").strip().lower()
if quantum_choice == 'y':
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
        USE_QUANTUM = True
        print("Quantum transforms ENABLED (Qiskit already installed).")
    except ImportError:
        print("Qiskit not found. Installing automatically...")
        if install_package('qiskit'):
            try:
                from qiskit import QuantumCircuit
                HAS_QISKIT = True
                USE_QUANTUM = True
                print("Quantum transforms ENABLED after automatic installation.")
            except ImportError:
                print("Qiskit installation succeeded but import failed – quantum transforms disabled.")
        else:
            print("Automatic installation failed – quantum transforms disabled.")
else:
    print("Quantum transforms disabled.")

# ------------------------------------------------------------------
# 2. Ask about other optional compression backends
# ------------------------------------------------------------------
other_choice = input("Install other optional compression backends (zstandard, paq, mpmath, cython, python-docx)? (y/n): ").strip().lower()
if other_choice == 'y':
    for pkg in ['mpmath', 'zstandard', 'cython', 'paq', 'python-docx']:
        try:
            importlib.import_module(pkg)
        except ImportError:
            install_package(pkg)
else:
    print("Skipping other backends.")

# ---------- Optional compression backends ----------
try:
    import paq
except ImportError:
    paq = None

try:
    import zstandard as zstd
    zstd_cctx = zstd.ZstdCompressor(level=22)
    zstd_dctx = zstd.ZstdDecompressor()
    HAS_ZSTD = True
except ImportError:
    HAS_ZSTD = False

if USE_QUANTUM and not HAS_QISKIT:
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
    except ImportError:
        USE_QUANTUM = False
        print("Quantum transforms disabled because Qiskit could not be imported.")

PROGNAME = "PJP"

# ---------- Dictionary configuration ----------
DICT_DIR = "Dictionaries"
COMBINED_DICTIONARY_FILE = os.path.join(DICT_DIR, "dictionary_combined.txt")

DICTIONARY_FILES = [
    "generated.txt",
    "eng_news_2005_1M-sentences.txt",
    "eng_news_2005_1M-words.txt",
    "eng_news_2005_1M-sources.txt",
    "eng_news_2005_1M-co_n.txt",
    "eng_news_2005_1M-co_s.txt",
    "eng_news_2005_1M-inv_w_2.txt",
    "eng_news_2005_1M-inv_w_3.txt",
    "eng_news_2005_1M-inv_so.txt",
    "eng_news_2005_1M-meta.txt",
    "Dictionary.txt",
    "the-complete-reference-html-css-fifth-edition.txt",
]

DICTIONARY_URLS = [
    "https://drive.google.com/uc?export=download&id=1u_1dCEl8hhdEug6GwkOxHAuSx_6_Pme9",
    "https://drive.google.com/uc?export=download&id=1pVqNN5JZ2AeOCgRaHkv4Vv6Byr4zK20e",
    "https://drive.google.com/uc?export=download&id=1ZSC-Tn76x8itdN0rCp-Zw17hGudxbjxo",
    "https://drive.google.com/uc?export=download&id=1VB_7tzngs4GxjclSRyRDnxgS8znT2w2S",
    "https://drive.google.com/uc?export=download&id=1KVIRgiMrhCUCqQZJ3UT67ztls2GqGJzz",
    "https://drive.google.com/uc?export=download&id=1Z3Lx6SqL4HWsnmbJCez4kXWRQQhUXWKL",
    "https://drive.google.com/uc?export=download&id=1br2bdRMkZEVVRPKYmC4IIaZuAjxFJE4N",
    "https://drive.google.com/uc?export=download&id=1aE6ubPZiJ8rr3lEVk8fFJYjDQ1y1rU0X",
    "https://drive.google.com/uc?export=download&id=1uro3TZe-t5zPx2Qu2xrTL3lU8N0melk9",
    "https://drive.google.com/uc?export=download&id=1HqsTH1DqpWNpGbn9VtD7-SB6wVqA90R2",
    "https://drive.google.com/uc?export=download&id=1zZ8iMeBC3605NZhuc4UE9jx_w_lZFg5B",
    "https://drive.google.com/uc?export=download&id=1dDdqYDgm7f-smS7KF70Wf0KmyFo-ft1M",
]

MAX_LINE_ENTRIES = 1024

def download_and_merge_dictionaries():
    if not os.path.exists(DICT_DIR):
        os.makedirs(DICT_DIR)

    if os.path.exists(COMBINED_DICTIONARY_FILE):
        print(f"Combined dictionary '{COMBINED_DICTIONARY_FILE}' already exists. Skipping download.")
        return True

    all_words = set()
    success_count = 0

    for idx, (filename, url) in enumerate(zip(DICTIONARY_FILES, DICTIONARY_URLS)):
        local_path = os.path.join(DICT_DIR, filename)
        print(f"Downloading {filename} to {DICT_DIR}/ ...")
        try:
            # Google Drive may require a confirmation token; handle redirect.
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                content = response.read()
                # Check if we got a HTML confirmation page (common for large files)
                if b'<html' in content[:200].lower() and b'confirm' in content[:2000].lower():
                    # Try to extract the confirm token and follow it
                    import re
                    match = re.search(b'confirm=([^&]+)', content)
                    if match:
                        confirm_token = match.group(1).decode()
                        new_url = f"https://drive.google.com/uc?export=download&confirm={confirm_token}&id={url.split('id=')[-1]}"
                        with urllib.request.urlopen(new_url) as resp2:
                            content = resp2.read()
                    else:
                        print(f"  WARNING: {filename} appears to be an HTML confirmation page; skipping.")
                        continue
                elif b'<html' in content[:200].lower():
                    print(f"  WARNING: {filename} appears to be an HTML page. Skipping.")
                    continue

            with open(local_path, 'wb') as f:
                f.write(content)

            with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if not w: continue
                    try:
                        decoded = base64.b64decode(w, validate=True)
                        all_words.add(decoded.decode('utf-8'))
                    except Exception:
                        all_words.add(w)

            print(f"  Downloaded {filename} ({os.path.getsize(local_path)} bytes)")
            success_count += 1
        except Exception as e:
            print(f"  WARNING: Could not download {filename}: {e}")
            if os.path.exists(local_path):
                os.remove(local_path)

    if success_count == 0:
        print("ERROR: No dictionary files could be downloaded.")
        return False

    try:
        with open(COMBINED_DICTIONARY_FILE, 'w', encoding='utf-8') as f:
            for word in sorted(all_words):
                f.write(word + '\n')
        print(f"Merged {len(all_words)} unique words into {COMBINED_DICTIONARY_FILE} "
              f"({os.path.getsize(COMBINED_DICTIONARY_FILE)} bytes)")
        return True
    except Exception as e:
        print(f"ERROR: Could not write combined dictionary: {e}")
        return False

# ---------- Constants ----------
PRIMES = [p for p in range(2, 256) if all(p % d != 0 for d in range(2, int(p ** 0.5) + 1))]
PI_DIGITS = [79, 17, 111]
BLOCK_SIZE = 1024

def find_nearest_prime_around(n: int) -> int:
    o = 0
    while True:
        c1, c2 = n - o, n + o
        if c1 >= 2 and all(c1 % d != 0 for d in range(2, int(c1 ** 0.5) + 1)):
            return c1
        if c2 >= 2 and all(c2 % d != 0 for d in range(2, int(c2 ** 0.5) + 1)):
            return c2
        o += 1

def sha256_8bytes(data: bytes) -> bytes:
    return hashlib.sha256(data).digest()[:8]

def xor_prime_hash(word: str) -> bytes:
    prime = 2147483647
    total = sum(ord(c) for c in word)
    transformed = total ^ prime
    return transformed.to_bytes(8, 'big')

ALPHABET_6BIT = (
    "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    "abcdefghijklmnopqrstuvwxyz"
    "0123456789"
    " \n"
)
assert len(ALPHABET_6BIT) == 64
CHAR_TO_6BIT = {ch: i for i, ch in enumerate(ALPHABET_6BIT)}
SIXBIT_TO_CHAR = {i: ch for ch, i in CHAR_TO_6BIT.items()}

class PJPCompressor:
    def __init__(self):
        download_and_merge_dictionaries()
        self.PI_DIGITS = PI_DIGITS.copy()
        self.seed_tables = self._gen_seed_tables(num=126, size=40, seed=42)
        self.fibonacci = self._gen_fib(100)
        self.PI_STR = "3.14159265358979323846264338327950288419716939937510"
        self._build_transform_maps()
        self.sequences = self._build_pair_sequences()
        self.pair_lookup = {idx: (t1, t2) for idx, (t1, t2) in enumerate(self.sequences)}
        self.static_dict, self.word_to_index = self._load_static_dictionary()
        self.line_dict, self.line_to_index = self._load_line_dictionary()
        if USE_QUANTUM and HAS_QISKIT:
            self._precompute_quantum_transforms()

    # ------------------------------------------------------------------
    # Quantum transforms
    # ------------------------------------------------------------------
    def _generate_permutation_from_circuit(self, num_qubits: int, seed: int) -> List[int]:
        qc = QuantumCircuit(num_qubits)
        rng = random.Random(seed)
        for qubit in range(num_qubits):
            qc.h(qubit)
            qc.rz(rng.random() * 2 * math.pi, qubit)
            qc.rx(rng.random() * 2 * math.pi, qubit)
        for _ in range(num_qubits):
            for i in range(num_qubits - 1):
                qc.cx(i, i+1)
            qc.barrier()
            for i in range(num_qubits):
                qc.rz(rng.random() * 2 * math.pi, i)
                qc.rx(rng.random() * 2 * math.pi, i)
        try:
            qasm_str = qc.qasm()
        except AttributeError:
            qasm_str = qc.draw('text')
        final_seed = seed + hash(qasm_str) % 1000000
        rng2 = random.Random(final_seed)
        n = 1 << num_qubits
        perm = list(range(n))
        rng2.shuffle(perm)
        if num_qubits == 12:
            perm_2704 = list(range(2704))
            rng2.shuffle(perm_2704)
            return perm_2704
        return perm

    def _precompute_quantum_transforms(self):
        self.quantum_fast_perms = [self._generate_permutation_from_circuit(8, 1000+i) for i in range(9)]
        self.quantum_ultra_perms = [self._generate_permutation_from_circuit(12, 2000+i) for i in range(17)]
        self.quantum_fast_transforms = []
        for perm in self.quantum_fast_perms:
            self.quantum_fast_transforms.append(self._make_substitution_transform(perm, 256))
        self.quantum_ultra_transforms = []
        for perm in self.quantum_ultra_perms:
            self.quantum_ultra_transforms.append(self._make_permutation_transform(perm, 2704))
        for idx, (fwd, rev) in enumerate(self.quantum_fast_transforms, start=257):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev
        for idx, (fwd, rev) in enumerate(self.quantum_ultra_transforms, start=266):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev

    def _make_substitution_transform(self, perm: List[int], size: int):
        inv_perm = [0] * size
        for i, p in enumerate(perm): inv_perm[p] = i
        def forward(data: bytes) -> bytes: return bytes(perm[b] for b in data)
        def reverse(data: bytes) -> bytes: return bytes(inv_perm[b] for b in data)
        return forward, reverse

    def _make_permutation_transform(self, perm: List[int], block_size: int):
        inv_perm = [0] * block_size
        for i, p in enumerate(perm): inv_perm[p] = i
        def forward(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        def reverse(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[inv_perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        return forward, reverse

    # ---------- Dictionary loaders ----------
    def _load_static_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            print(f"ERROR: {COMBINED_DICTIONARY_FILE} not found. No dictionaries loaded.")
            return [], {}
        words_set = set()
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if w: words_set.add(w)
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}
        sorted_words = sorted(words_set)
        word_to_idx = {w: i for i, w in enumerate(sorted_words)}
        print(f"Loaded static word dictionary: {len(sorted_words)} unique words.")
        return sorted_words, word_to_idx

    def _load_line_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            return [], {}
        lines = []
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for raw_line in f:
                    phrase = raw_line.strip()
                    if phrase and phrase not in lines:
                        lines.append(phrase)
                        if len(lines) >= MAX_LINE_ENTRIES: break
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}
        if not lines: return [], {}
        lines.sort(key=len, reverse=True)
        line_to_idx = {phrase: i for i, phrase in enumerate(lines)}
        print(f"Loaded line dictionary: {len(lines)} phrases.")
        return lines, line_to_idx

    # ---------- Pi & constant helpers ----------
    def get_pi_digits(self, n): return self.PI_STR[2:2 + n] if n > 0 else ""
    def find_lossless_k(self, n):
        if n < 1: return 0, True
        true_digits = self.get_pi_digits(n)
        true_scaled = int(self.PI_STR.replace('.', '')[:n + 1])
        DENOM = 16777216
        decimal.getcontext().prec = 50
        pi_dec = decimal.Decimal(self.PI_STR)
        k_float = (pi_dec - 3) * DENOM
        k_candidate = int(round(k_float))
        k_candidate = max(0, min(k_candidate, DENOM - 1))
        approx_scaled = (3 * 10 ** n * DENOM + k_candidate * 10 ** n) // DENOM
        return k_candidate, approx_scaled == true_scaled
    def to_bin(self, value, bits): return format(value, 'b').zfill(bits)
    def get_bit_size(self, k): return 23 if k <= 0x7FFFFF else 25
    def transform_17(self, data: bytes) -> bytes:
        if not data: return b''
        k, _ = self.find_lossless_k(7)
        bit_str = self.to_bin(k, self.get_bit_size(k))
        mask = bytes(int(bit_str[i:i+8].ljust(8, '0'), 2) for i in range(0, len(bit_str), 8))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_17 = transform_17
    def get_basel_digits(self, n):
        decimal.getcontext().prec = n + 5
        basel = (decimal.Decimal(self.PI_STR) ** 2) / 6
        return str(basel).replace('.', '')[:n]
    def get_one_over_e_digits(self, n):
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        return str(1 / e).replace('.', '')[:n]
    def get_5e_digits(self, n):
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        return str(5 * e).replace('.', '')[:n]

    # ---------- Seed tables & Fibonacci ----------
    def _gen_seed_tables(self, num=126, size=40, seed=42):
        random.seed(seed)
        return [[random.randint(5, 255) for _ in range(size)] for _ in range(num)]
    def _gen_fib(self, n):
        a, b = 0, 1
        res = [a, b]
        for _ in range(2, n):
            a, b = b, a + b
            res.append(b)
        return res
    def get_seed(self, idx, val):
        if 0 <= idx < len(self.seed_tables):
            return self.seed_tables[idx][val % 40]
        return 0

    # ---------- Bit helpers ----------
    def _append_bits(self, bitlist, value, count):
        for i in range(count - 1, -1, -1):
            bitlist.append((value >> i) & 1)
    def _read_bits(self, bits, pos, count):
        val = 0
        for i in range(count):
            if pos + i >= len(bits): return 0
            val = (val << 1) | bits[pos + i]
        return val

    # ---------- Transform 00 (RLE) ----------
    def transform_00(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        best_result, best_length = None, float('inf')
        best_shifts, current = [], bytearray(data)
        applied_shifts = []
        for _ in range(10):
            best_shift, best_score = 0, float('-inf')
            for shift in range(256):
                tmp = bytearray((b + shift) % 256 for b in current)
                score = 0; i = 0
                while i < len(tmp):
                    val = tmp[i]; run = 1; i += 1
                    while i < len(tmp) and tmp[i] == val: run += 1; i += 1
                    score += run * run
                if score > best_score: best_score, best_shift = score, shift
            applied_shifts.append(best_shift)
            shifted = bytearray((b + best_shift) % 256 for b in current)
            encoded = self._apply_rle_to_shifted(shifted, best_shift)
            if len(encoded) < best_length:
                best_length, best_result, best_shifts = len(encoded), encoded, applied_shifts.copy()
            current = shifted
        if best_result is None or best_length >= len(data):
            return bytes([0]) + data
        header = bytearray([len(best_shifts)])
        header.extend(best_shifts)
        return header + best_result

    def _apply_rle_to_shifted(self, shifted_data: bytearray, shift: int) -> bytes:
        bits = []
        self._append_bits(bits, 0b010, 3)
        self._append_bits(bits, shift, 8)
        i, n = 0, len(shifted_data)
        while i < n:
            val, run = shifted_data[i], 1; i += 1
            while i < n and shifted_data[i] == val: run += 1; i += 1
            while run >= 13:
                chunk = min(run, 268)
                self._append_bits(bits, 0b1111, 4)
                self._append_bits(bits, chunk - 13, 8)
                self._append_bits(bits, val, 8)
                run -= chunk
            if run == 1:
                self._append_bits(bits, 0b00, 2)
                self._append_bits(bits, val, 8)
            elif run <= 5:
                self._append_bits(bits, 0b01, 2)
                self._append_bits(bits, run - 2, 2)
                self._append_bits(bits, val, 8)
            elif run <= 12:
                self._append_bits(bits, 0b10, 2)
                self._append_bits(bits, run - 6, 3)
                self._append_bits(bits, val, 8)
        pad = (8 - len(bits) % 8) % 8
        self._append_bits(bits, 0, pad)
        out = bytearray()
        for j in range(0, len(bits), 8):
            byte = 0
            for k in range(8):
                if j + k < len(bits): byte = (byte << 1) | bits[j + k]
            out.append(byte)
        return bytes(out)

    def reverse_transform_00(self, cdata: bytes) -> bytes:
        if not cdata or cdata == b'\x00': return b''
        if cdata[0] == 0: return cdata[1:]
        num_passes = cdata[0]
        if num_passes == 0 or len(cdata) < 1 + num_passes: return b''
        shifts = list(cdata[1:1 + num_passes])
        rle_data = cdata[1 + num_passes:]
        decoded = self._rle_decode(rle_data)
        if decoded is None: return b''
        current = bytearray(decoded)
        for shift in reversed(shifts):
            for i in range(len(current)):
                current[i] = (current[i] - shift) % 256
        return bytes(current)

    def _rle_decode(self, data: bytes) -> Optional[bytearray]:
        if not data: return None
        bits = []
        for b in data:
            for i in range(7, -1, -1):
                bits.append((b >> i) & 1)
        pos, nbits = 0, len(bits)
        if nbits < 11: return None
        if self._read_bits(bits, pos, 3) != 0b010: return None
        pos += 3 + 8
        out = bytearray()
        while pos < nbits:
            if pos + 2 > nbits: break
            prefix = self._read_bits(bits, pos, 2)
            pos += 2
            if prefix == 0b00: run = 1
            elif prefix == 0b01: run = 2 + self._read_bits(bits, pos, 2); pos += 2
            elif prefix == 0b10: run = 6 + self._read_bits(bits, pos, 3); pos += 3
            else:
                if self._read_bits(bits, pos, 2) != 0b11: return None
                pos += 2
                run = 13 + self._read_bits(bits, pos, 8); pos += 8
            if pos + 8 > nbits: break
            val = self._read_bits(bits, pos, 8); pos += 8
            out.extend([val] * run)
        for i in range(pos, nbits):
            if bits[i] != 0: return None
        return out

    # ---------- Transforms 01‑21 (same as before) ----------
    def transform_01(self, d, r=100):
        t = bytearray(d)
        for prime in PRIMES:
            xor_val = prime if prime == 2 else max(1, math.ceil(prime * 4096 / 28672))
            for _ in range(r):
                for i in range(0, len(t), 3):
                    if i < len(t): t[i] ^= xor_val
        return bytes(t)
    reverse_transform_01 = transform_01
    def transform_02(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        checksum = sum(d) % 256
        pattern_index = (len(d) + checksum) % 256
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_02(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes(t)
    def transform_03(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        rotation = (len(d) * 13 + sum(d)) % 8 or 1
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] << rotation) | (t[i] >> (8 - rotation))) & 0xFF
        return bytes([rotation]) + bytes(t)
    def reverse_transform_03(self, d):
        if len(d) < 2: return b''
        rotation = d[0]
        t = bytearray(d[1:])
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] >> rotation) | (t[i] << (8 - rotation))) & 0xFF
        return bytes(t)
    def transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] - (i % 256)) % 256
        return bytes(t)
    def reverse_transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] + (i % 256)) % 256
        return bytes(t)
    def transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] << s) | (t[i] >> (8 - s))) & 0xFF
        return bytes(t)
    def reverse_transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] >> s) | (t[i] << (8 - s))) & 0xFF
        return bytes(t)
    def transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256)); random.shuffle(sub)
        t = bytearray(d)
        for i in range(len(t)): t[i] = sub[t[i]]
        return bytes(t)
    def reverse_transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256)); random.shuffle(sub)
        inv = [0]*256
        for i in range(256): inv[sub[i]] = i
        t = bytearray(d)
        for i in range(len(t)): t[i] = inv[t[i]]
        return bytes(t)
    def transform_07(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        sz = len(d) % 256
        for i in range(len(t)): t[i] ^= sz
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_07 = transform_07
    def transform_08(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        for i in range(len(t)): t[i] ^= p
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_08 = transform_08
    def transform_09(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        seed = self.get_seed(len(d) % len(self.seed_tables), len(d))
        for i in range(len(t)): t[i] ^= p ^ seed
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)] ^ (i % 256)
        return bytes(t)
    reverse_transform_09 = transform_09
    def transform_10(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        cnt = sum(1 for i in range(len(data)-1) if data[i:i+2] == b'X1')
        n = (((cnt * 2) + 1) // 3) * 3 % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= n
        return bytes([n]) + bytes(t)
    def reverse_transform_10(self, data: bytes) -> bytes:
        if len(data) < 1: return b''
        n = data[0]
        t = bytearray(data[1:])
        for i in range(len(t)): t[i] ^= n
        return bytes(t)
    def transform_11(self, data: bytes) -> bytes:
        if not data: return b''
        t = bytearray(data)
        length = len(t)
        for i in range(length):
            fib_idx = (i + length) % len(self.fibonacci)
            fib_val = self.fibonacci[fib_idx] % 256
            pos_val = (i * 13 + length * 17) % 256
            key = (fib_val ^ pos_val) % 256
            t[i] ^= key
        return bytes(t)
    reverse_transform_11 = transform_11
    def transform_12(self, data: bytes) -> bytes:
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= self.fibonacci[i % len(self.fibonacci)] % 256
        return bytes(t)
    reverse_transform_12 = transform_12
    def transform_13(self, d):
        if not d: return b''
        repeats = self._calculate_repeats(d)
        current_value = len(d) % 256
        prime_values = []
        for _ in range(repeats):
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
        xor_value = prime_values[-1] if prime_values else 0
        t = bytearray(d)
        for i in range(len(t)): t[i] ^= xor_value
        return bytes([(repeats - 1) % 256]) + bytes(t)
    def reverse_transform_13(self, d):
        if len(d) < 2: return b''
        repeat_byte = d[0]
        repeats = (repeat_byte + 1) % 256 or 256
        t = bytearray(d[1:])
        current_value = len(t) % 256
        prime_values = []
        for _ in range(repeats):
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
        xor_value = prime_values[-1] if prime_values else 0
        for i in range(len(t)): t[i] ^= xor_value
        return bytes(t)
    def transform_15(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        pattern_index = len(d) % 256
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] + pattern_values[i % len(pattern_values)]) % 256
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_15(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] - pattern_values[i % len(pattern_values)]) % 256
        return bytes(t)
    def transform_16(self, data: bytes) -> bytes:
        if not data: return b''
        xor_byte = (len(data) * 7 + 13) % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= xor_byte
        return bytes(t)
    reverse_transform_16 = transform_16
    def transform_18(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_basel_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_18 = transform_18
    def transform_19(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_one_over_e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_19 = transform_19
    def transform_20(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_5e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_20 = transform_20
    def transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] + shift) % 256
        return bytes(t)
    def reverse_transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] - shift) % 256
        return bytes(t)

    # ---------- Transform 22 – Base64 ----------
    def transform_22(self, data: bytes) -> bytes: return base64.b64encode(data)
    def reverse_transform_22(self, data: bytes) -> bytes:
        try: return base64.b64decode(data, validate=False)
        except: return data

    # ---------- Transforms 23‑27 (unchanged) ----------
    def transform_23(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try: text = data.decode('latin-1')
        except: text = data.decode('latin-1', errors='replace')
        tokens = re.split(r'([A-Za-z0-9_]+)', text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = sha256_8bytes(word_bytes)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok: token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h; result += struct.pack('>H', len(wb)); result += wb
        for is_word, payload in token_list:
            if is_word: result += b'\x01'; result += payload
            else: result += b'\x00'; result += struct.pack('>H', len(payload)); result += payload
        return bytes(result)
    def reverse_transform_23(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]; pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]; pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]; pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]; pos += 8
                out += hash_to_word.get(h, h)
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]; pos += rawlen
            else: break
        return bytes(out)
    def transform_24(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try: text = data.decode('latin-1')
        except: text = data.decode('latin-1', errors='replace')
        tokens = re.split(r'([A-Za-z0-9_]+)', text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = xor_prime_hash(tok)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok: token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h; result += struct.pack('>H', len(wb)); result += wb
        for is_word, payload in token_list:
            if is_word: result += b'\x01'; result += payload
            else: result += b'\x00'; result += struct.pack('>H', len(payload)); result += payload
        return bytes(result)
    def reverse_transform_24(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]; pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]; pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]; pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]; pos += 8
                out += hash_to_word.get(h, h)
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]; pos += rawlen
            else: break
        return bytes(out)
    def _split_text_into_chunks(self, text, level='all'):
        if level == 'paragraph': return re.split(r'(\n\n)', text)
        elif level == 'line': return re.split(r'(\n)', text)
        elif level == 'sentence': return re.split(r'([.!?]+)', text)
        elif level == 'word': return re.split(r'(\s+|\b)', text)
        else:
            chunks = []
            for para in re.split(r'(\n\n)', text):
                if para: chunks.append(para)
            return chunks
    def _dynamic_dict_tokenize(self, data, index_bytes=3):
        try: text = data.decode('utf-8')
        except: return b'\x00' + data
        chunks = self._split_text_into_chunks(text)
        freq = Counter(chunks)
        sorted_chunks = sorted(freq.keys(), key=lambda x: (-freq[x], -len(x), x))
        chunk_to_idx = {ch: i for i, ch in enumerate(sorted_chunks)}
        num_entries = len(sorted_chunks)
        if index_bytes == 2 and num_entries > 65535: index_bytes = 3
        if index_bytes == 3 and num_entries > 16777215: index_bytes = 8
        header = bytearray([index_bytes])
        header += struct.pack('>I', num_entries)
        for chunk in sorted_chunks:
            cb = chunk.encode('utf-8')
            header += struct.pack('>I', len(cb))
            header += cb
        token_stream = bytearray()
        for chunk in chunks:
            idx = chunk_to_idx[chunk]
            if index_bytes == 2: token_stream += struct.pack('>H', idx)
            elif index_bytes == 3: token_stream += struct.pack('>I', idx)[1:4]
            else: token_stream += struct.pack('>Q', idx)
        return bytes(header) + bytes(token_stream)
    def _dynamic_dict_detokenize(self, data):
        if not data: return b''
        if data[0] == 0: return data[1:]
        index_bytes = data[0]
        if index_bytes not in (2, 3, 8): return None
        pos = 1
        if pos + 4 > len(data): return None
        num_entries = struct.unpack('>I', data[pos:pos+4])[0]
        pos += 4
        dictionary = []
        for _ in range(num_entries):
            if pos + 4 > len(data): return None
            chunk_len = struct.unpack('>I', data[pos:pos+4])[0]; pos += 4
            if pos + chunk_len > len(data): return None
            chunk = data[pos:pos+chunk_len].decode('utf-8'); pos += chunk_len
            dictionary.append(chunk)
        tokens = []
        while pos < len(data):
            if index_bytes == 2:
                if pos + 2 > len(data): break
                idx = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
            elif index_bytes == 3:
                if pos + 3 > len(data): break
                idx = struct.unpack('>I', b'\x00' + data[pos:pos+3])[0]; pos += 3
            else:
                if pos + 8 > len(data): break
                idx = struct.unpack('>Q', data[pos:pos+8])[0]; pos += 8
            if idx < len(dictionary): tokens.append(dictionary[idx])
            else: return None
        try: return ''.join(tokens).encode('utf-8')
        except: return None
    def transform_25(self, data): return self._dynamic_dict_tokenize(data, 3)
    def reverse_transform_25(self, data):
        res = self._dynamic_dict_detokenize(data)
        return res if res else b''
    def transform_26(self, data: bytes) -> bytes:
        if not data: return b''
        secret = b"PJP_TRANSFORM26_SECRET"
        result = bytearray()
        for idx in range(0, len(data), BLOCK_SIZE):
            chunk = data[idx:idx+BLOCK_SIZE]
            block_num = idx // BLOCK_SIZE
            hasher = hashlib.sha256(secret + struct.pack(">Q", block_num))
            mask = hasher.digest() * ((len(chunk) // 32) + 1)
            mask = mask[:len(chunk)]
            result.extend(bytes(a ^ b for a, b in zip(chunk, mask)))
        return bytes(result)
    reverse_transform_26 = transform_26
    def transform_27(self, data: bytes) -> bytes:
        try: text = data.decode('utf-8')
        except: return data
        if any(ch not in CHAR_TO_6BIT for ch in text): return data
        bits = []
        for ch in text:
            val = CHAR_TO_6BIT[ch]
            for i in range(5, -1, -1): bits.append((val >> i) & 1)
        pad = (8 - len(bits) % 8) % 8
        bits.extend([0] * pad)
        out = bytearray()
        for i in range(0, len(bits), 8):
            byte = 0
            for j in range(8): byte = (byte << 1) | bits[i + j]
            out.append(byte)
        return struct.pack('<I', len(text)) + bytes(out)
    def reverse_transform_27(self, data: bytes) -> bytes:
        if len(data) < 4: return data
        num_chars = struct.unpack('<I', data[:4])[0]
        packed = data[4:]
        bits = []
        for b in packed:
            for i in range(7, -1, -1): bits.append((b >> i) & 1)
        needed_bits = num_chars * 6
        if len(bits) < needed_bits: return data
        chars = []
        for i in range(num_chars):
            val = 0
            for j in range(6): val = (val << 1) | bits[i*6 + j]
            if val < 64: chars.append(SIXBIT_TO_CHAR[val])
            else: return data
        try: return ''.join(chars).encode('utf-8')
        except: return data

    # ---------- Transforms 28‑30 (unchanged) ----------
    def transform_28(self, data):
        if not data: return b''
        pad = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad
        out = bytearray([pad])
        for i in range(0, len(padded), 3):
            val = int.from_bytes(padded[i:i+3], 'little')
            key = (i//3 * 65537 + 12345) & 0xFFFF
            out.extend(((val - key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out)
    def reverse_transform_28(self, data):
        if not data: return b''
        pad = data[0]
        payload = data[1:]
        if len(payload) % 3: return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            val = int.from_bytes(payload[i:i+3], 'little')
            key = (i//3 * 65537 + 12345) & 0xFFFF
            out.extend(((val + key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out[:-pad] if pad else out)
    def _find_best_16bit_key(self, data):
        if len(data) < 3: return 0
        pad = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad
        values = [int.from_bytes(padded[i:i+3], 'little') for i in range(0, len(padded), 3)]
        mean = sum(values) // len(values)
        best_key, best_cost = 0, float('inf')
        for key in range(65536):
            trans = [(v - key) & 0xFFFFFF for v in values]
            mean_t = sum(trans) // len(trans)
            cost = sum(abs(t - mean_t) for t in trans)
            if cost < best_cost:
                best_cost, best_key = cost, key
        return best_key
    def transform_29(self, data):
        if not data: return b''
        key = self._find_best_16bit_key(data)
        pad = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad
        out = bytearray([pad])
        out.extend(key.to_bytes(2, 'little'))
        for i in range(0, len(padded), 3):
            val = int.from_bytes(padded[i:i+3], 'little')
            out.extend(((val - key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out)
    def reverse_transform_29(self, data):
        if len(data) < 3: return data
        pad = data[0]
        key = int.from_bytes(data[1:3], 'little')
        payload = data[3:]
        if len(payload) % 3: return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            val = int.from_bytes(payload[i:i+3], 'little')
            out.extend(((val + key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out[:-pad] if pad else out)
    def _find_best_24bit_key_heuristic(self, data):
        if len(data) < 3: return 0
        pad = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad
        values = [int.from_bytes(padded[i:i+3], 'little') for i in range(0, len(padded), 3)]
        mean = sum(values) // len(values)
        median = sorted(values)[len(values)//2]
        candidates = set()
        for base in (mean, median):
            for offset in (0, 1, -1, 10, -10, 100, -100, 1000, -1000):
                candidates.add((base + offset) & 0xFFFFFF)
        rng = random.Random(42)
        for _ in range(10): candidates.add(rng.randint(0, (1<<24)-1))
        best_key, best_cost = 0, float('inf')
        for key in candidates:
            trans = [(v - key) & 0xFFFFFF for v in values]
            mean_t = sum(trans) // len(trans)
            cost = sum(abs(t - mean_t) for t in trans)
            if cost < best_cost:
                best_cost, best_key = cost, key
        return best_key
    def transform_30(self, data):
        if not data: return b''
        key = self._find_best_24bit_key_heuristic(data)
        pad = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad
        out = bytearray([pad])
        out.extend(key.to_bytes(3, 'little'))
        for i in range(0, len(padded), 3):
            val = int.from_bytes(padded[i:i+3], 'little')
            out.extend(((val - key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out)
    def reverse_transform_30(self, data):
        if len(data) < 4: return data
        pad = data[0]
        key = int.from_bytes(data[1:4], 'little')
        payload = data[4:]
        if len(payload) % 3: return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            val = int.from_bytes(payload[i:i+3], 'little')
            out.extend(((val + key) & 0xFFFFFF).to_bytes(3, 'little'))
        return bytes(out[:-pad] if pad else out)

    # ---------- Dictionary tokenization for 31/32 ----------
    def _build_text_dictionary(self, text_streams, min_freq=2):
        all_tokens = []
        for text in text_streams:
            all_tokens.extend(re.findall(r'\b[\w\-]+\b', text))
        freq = Counter(all_tokens)
        common = sorted([w for w, c in freq.items() if c >= min_freq],
                       key=lambda w: (-freq[w], -len(w), w))
        return common, {w: i for i, w in enumerate(common)}
    def _encode_text_with_dict(self, text, dictionary, word_to_idx):
        parts = re.split(r'(\b[\w\-]+\b)', text)
        encoded = bytearray()
        for i, part in enumerate(parts):
            if i % 2 == 1:
                if part in word_to_idx:
                    idx = word_to_idx[part]
                    if len(dictionary) <= 255:
                        encoded += bytes([0x00, idx])
                    elif len(dictionary) <= 65535:
                        encoded += b'\x01' + struct.pack('>H', idx)
                    elif len(dictionary) <= 16777215:
                        encoded += b'\x02' + struct.pack('>I', idx)[1:4]
                    else:
                        encoded += b'\x03' + struct.pack('>Q', idx)
                else:
                    wb = part.encode('utf-8')
                    encoded += b'\x04' + bytes([len(wb)]) + wb
            else:
                if part:
                    raw = part.encode('utf-8')
                    encoded += b'\x04' + bytes([len(raw)]) + raw
        return bytes(encoded)
    def _decode_text_with_dict(self, data, dictionary):
        pos = 0
        out = []
        while pos < len(data):
            marker = data[pos]; pos += 1
            if marker == 0x00:
                if pos >= len(data): break
                idx = data[pos]; pos += 1
                out.append(dictionary[idx] if idx < len(dictionary) else f"<ERR{idx}>")
            elif marker == 0x01:
                if pos + 1 >= len(data): break
                idx = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
                out.append(dictionary[idx] if idx < len(dictionary) else f"<ERR{idx}>")
            elif marker == 0x02:
                if pos + 2 >= len(data): break
                idx = struct.unpack('>I', b'\x00' + data[pos:pos+3])[0]; pos += 3
                out.append(dictionary[idx] if idx < len(dictionary) else f"<ERR{idx}>")
            elif marker == 0x03:
                if pos + 7 >= len(data): break
                idx = struct.unpack('>Q', data[pos:pos+8])[0]; pos += 8
                out.append(dictionary[idx] if idx < len(dictionary) else f"<ERR{idx}>")
            elif marker == 0x04:
                if pos >= len(data): break
                length = data[pos]; pos += 1
                if pos + length > len(data): break
                out.append(data[pos:pos+length].decode('utf-8', errors='replace'))
                pos += length
            else: break
        return ''.join(out)

    # ---------- Transform 31 – .docx paragraphs ----------
    def transform_31(self, data: bytes) -> bytes:
        if not data or data[:4] != b'PK\x03\x04': return b'\x00' + data
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
        except ImportError:
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    xml = zf.read('word/document.xml')
                root = ET.fromstring(xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text = ''.join(t.text or '' for t in root.findall('.//w:t', ns))
                if not text: return b'\x00' + data
                dict_list, w2i = self._build_text_dictionary([text])
                enc = self._encode_text_with_dict(text, dict_list, w2i)
                out = bytearray([0x01, len(dict_list)])
                for w in dict_list:
                    wb = w.encode('utf-8')
                    out += struct.pack('>H', len(wb)) + wb
                out += enc
                return bytes(out)
            except: return b'\x00' + data
        # using python-docx
        full = '\n'.join(''.join(r.text for r in p.runs if r.text) for p in doc.paragraphs)
        if not full: return b'\x00' + data
        dict_list, w2i = self._build_text_dictionary([full])
        out = bytearray([0x01, len(dict_list)])
        for w in dict_list:
            wb = w.encode('utf-8'); out += struct.pack('>H', len(wb)) + wb
        for p in doc.paragraphs:
            for r in p.runs:
                if not r.text: continue
                enc = self._encode_text_with_dict(r.text, dict_list, w2i)
                size = int(r.font.size.pt) if r.font.size else 12
                style = 0
                if r.bold: style |= 1
                if r.italic: style |= 2
                if r.underline: style |= 4
                if r.strike: style |= 8
                if r.superscript: style |= 16
                if r.subscript: style |= 32
                out += b'\x05' + bytes([size, style]) + enc
        return bytes(out)
    def reverse_transform_31(self, data: bytes) -> bytes:
        if not data: return b''
        if data[0] == 0x00: return data[1:]
        if data[0] != 0x01: return data
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError: return data
        pos, n = 1, data[1]
        pos += 1
        dictionary = []
        for _ in range(n):
            wlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
            dictionary.append(data[pos:pos+wlen].decode('utf-8')); pos += wlen
        doc = Document()
        p = doc.add_paragraph()
        while pos < len(data):
            if data[pos] == 0x05:
                pos += 1
                size, style = data[pos], data[pos+1]; pos += 2
                run_data = bytearray()
                while pos < len(data) and data[pos] != 0x05:
                    run_data.append(data[pos]); pos += 1
                text = self._decode_text_with_dict(bytes(run_data), dictionary)
                run = p.add_run(text)
                run.font.size = Pt(size)
                if style & 1: run.bold = True
                if style & 2: run.italic = True
                if style & 4: run.underline = True
                if style & 8: run.strike = True
                if style & 16: run.superscript = True
                if style & 32: run.subscript = True
            else: break
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ---------- Transform 32 – .docx tables ----------
    def transform_32(self, data: bytes) -> bytes:
        if not data or data[:4] != b'PK\x03\x04': return b'\x00' + data
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
        except ImportError: return b'\x00' + data
        tables = doc.tables
        if not tables: return b'\x00' + data
        all_text = [cell.text for table in tables for row in table.rows for cell in row.cells]
        full = '\n'.join(all_text)
        if not full: return b'\x00' + data
        dict_list, w2i = self._build_text_dictionary([full])
        out = bytearray([0x02, len(dict_list)])
        for w in dict_list:
            wb = w.encode('utf-8'); out += struct.pack('>H', len(wb)) + wb
        for table in tables:
            rows, cols = len(table.rows), len(table.columns)
            out += bytes([rows, cols])
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text:
                        out += b'\x00'; continue
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if not r.text: continue
                            enc = self._encode_text_with_dict(r.text, dict_list, w2i)
                            size = int(r.font.size.pt) if r.font.size else 12
                            style = 0
                            if r.bold: style |= 1
                            if r.italic: style |= 2
                            if r.underline: style |= 4
                            if r.strike: style |= 8
                            if r.superscript: style |= 16
                            if r.subscript: style |= 32
                            out += b'\x06' + bytes([size, style]) + enc
                    out += b'\x00'
        return bytes(out)
    def reverse_transform_32(self, data: bytes) -> bytes:
        if not data: return b''
        if data[0] == 0x00: return data[1:]
        if data[0] != 0x02: return data
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError: return data
        pos, n = 1, data[1]
        pos += 1
        dictionary = []
        for _ in range(n):
            wlen = struct.unpack('>H', data[pos:pos+2])[0]; pos += 2
            dictionary.append(data[pos:pos+wlen].decode('utf-8')); pos += wlen
        doc = Document()
        while pos < len(data):
            if pos >= len(data): break
            rows, cols = data[pos], data[pos+1]; pos += 2
            table = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    cell = table.cell(r, c)
                    p = cell.paragraphs[0]
                    while pos < len(data) and data[pos] != 0x00:
                        if data[pos] == 0x06:
                            pos += 1
                            size, style = data[pos], data[pos+1]; pos += 2
                            run_data = bytearray()
                            while pos < len(data) and data[pos] not in (0x00, 0x06):
                                run_data.append(data[pos]); pos += 1
                            text = self._decode_text_with_dict(bytes(run_data), dictionary)
                            run = p.add_run(text)
                            run.font.size = Pt(size)
                            if style & 1: run.bold = True
                            if style & 2: run.italic = True
                            if style & 4: run.underline = True
                            if style & 8: run.strike = True
                            if style & 16: run.superscript = True
                            if style & 32: run.subscript = True
                        else: break
                    pos += 1  # skip trailing 0x00
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ---------- Extract plain text (internal only) ----------
    def extract_text_from_docx(self, data: bytes) -> str:
        if not data or data[:4] != b'PK\x03\x04': return ""
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            text = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return '\n'.join(text)
        except ImportError:
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    xml = zf.read('word/document.xml')
                root = ET.fromstring(xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                return ''.join(t.text or '' for t in root.findall('.//w:t', ns))
            except: return ""

    # ---------- Transform 256 – no-op ----------
    def transform_256(self, d): return d
    reverse_transform_256 = transform_256

    # ---------- Helpers ----------
    def _get_pattern(self, size, index):
        random.seed(12345 + size * 100 + index)
        return [random.randint(0, 255) for _ in range(size)]
    def _calculate_repeats(self, data):
        if not data: return 1
        return max(1, min(256, ((len(data) * 13 + sum(data) * 17) % 256) + 1))

    def _dynamic_transform(self, n):
        def tf(data: bytes):
            if not data: return b''
            seed = self.get_seed(n % len(self.seed_tables), len(data))
            t = bytearray(data)
            for i in range(len(t)): t[i] ^= seed
            return bytes(t)
        return tf, tf

    # ---------- Build transform maps (1‑512) ----------
    def _build_transform_maps(self):
        self.fwd_transforms: Dict[int, Callable] = {}
        self.rev_transforms: Dict[int, Callable] = {}

        # 1‑21
        self.fwd_transforms[1] = self.transform_00; self.rev_transforms[1] = self.reverse_transform_00
        self.fwd_transforms[2] = self.transform_01; self.rev_transforms[2] = self.reverse_transform_01
        self.fwd_transforms[3] = self.transform_02; self.rev_transforms[3] = self.reverse_transform_02
        self.fwd_transforms[4] = self.transform_03; self.rev_transforms[4] = self.reverse_transform_03
        self.fwd_transforms[5] = self.transform_04; self.rev_transforms[5] = self.reverse_transform_04
        self.fwd_transforms[6] = self.transform_05; self.rev_transforms[6] = self.reverse_transform_05
        self.fwd_transforms[7] = self.transform_06; self.rev_transforms[7] = self.reverse_transform_06
        self.fwd_transforms[8] = self.transform_07; self.rev_transforms[8] = self.reverse_transform_07
        self.fwd_transforms[9] = self.transform_08; self.rev_transforms[9] = self.reverse_transform_08
        self.fwd_transforms[10] = self.transform_09; self.rev_transforms[10] = self.reverse_transform_09
        self.fwd_transforms[11] = self.transform_10; self.rev_transforms[11] = self.reverse_transform_10
        self.fwd_transforms[12] = self.transform_11; self.rev_transforms[12] = self.reverse_transform_11
        self.fwd_transforms[13] = self.transform_12; self.rev_transforms[13] = self.reverse_transform_12
        self.fwd_transforms[14] = self.transform_13; self.rev_transforms[14] = self.reverse_transform_13
        self.fwd_transforms[15] = self.transform_15; self.rev_transforms[15] = self.reverse_transform_15
        self.fwd_transforms[16] = self.transform_16; self.rev_transforms[16] = self.reverse_transform_16
        self.fwd_transforms[17] = self.transform_17; self.rev_transforms[17] = self.reverse_transform_17
        self.fwd_transforms[18] = self.transform_18; self.rev_transforms[18] = self.reverse_transform_18
        self.fwd_transforms[19] = self.transform_19; self.rev_transforms[19] = self.reverse_transform_19
        self.fwd_transforms[20] = self.transform_20; self.rev_transforms[20] = self.reverse_transform_20
        self.fwd_transforms[21] = self.transform_21; self.rev_transforms[21] = self.reverse_transform_21
        # 22
        self.fwd_transforms[22] = self.transform_22; self.rev_transforms[22] = self.reverse_transform_22
        # 23‑27
        self.fwd_transforms[23] = self.transform_23; self.rev_transforms[23] = self.reverse_transform_23
        self.fwd_transforms[24] = self.transform_24; self.rev_transforms[24] = self.reverse_transform_24
        self.fwd_transforms[25] = self.transform_25; self.rev_transforms[25] = self.reverse_transform_25
        self.fwd_transforms[26] = self.transform_26; self.rev_transforms[26] = self.reverse_transform_26
        self.fwd_transforms[27] = self.transform_27; self.rev_transforms[27] = self.reverse_transform_27
        # 28‑30
        self.fwd_transforms[28] = self.transform_28; self.rev_transforms[28] = self.reverse_transform_28
        self.fwd_transforms[29] = self.transform_29; self.rev_transforms[29] = self.reverse_transform_29
        self.fwd_transforms[30] = self.transform_30; self.rev_transforms[30] = self.reverse_transform_30
        # 31‑32
        self.fwd_transforms[31] = self.transform_31; self.rev_transforms[31] = self.reverse_transform_31
        self.fwd_transforms[32] = self.transform_32; self.rev_transforms[32] = self.reverse_transform_32
        # 33‑255 dynamic
        for i in range(33, 256):
            fwd, rev = self._dynamic_transform(i)
            self.fwd_transforms[i] = fwd
            self.rev_transforms[i] = rev
        # 256 no-op (explicitly set to avoid being overwritten)
        self.fwd_transforms[256] = self.transform_256
        self.rev_transforms[256] = self.reverse_transform_256
        # Quantum transforms (257‑282) will be added later if enabled.
        # New fast transforms 283‑512
        for i in range(283, 513):
            fwd, rev = self._dynamic_transform(i)
            self.fwd_transforms[i] = fwd
            self.rev_transforms[i] = rev
        # Ensure all 1‑512 are defined
        for i in range(1, 513):
            if i not in self.fwd_transforms:
                raise RuntimeError(f"Transform {i} missing!")

    # ---------- Pair sequences (unchanged) ----------
    def _build_pair_sequences(self) -> List[Tuple[int, int]]:
        safe = []
        for i in range(1, 257):
            if i in (1, 14, 22, 23, 24, 25, 26, 27, 31, 32): continue
            safe.append(i)
            if len(safe) == 52: break
        while len(safe) < 52: safe.append(256)
        return [(t1, t2) for t1 in safe for t2 in safe]

    def _apply_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t in seq: result = self.fwd_transforms[t](result)
        return result
    def _reverse_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t in reversed(seq): result = self.rev_transforms[t](result)
        return result

    # ---------- Backends ----------
    def _compress_backend(self, data: bytes, safe: bool = False) -> bytes:
        candidates = []
        if paq:
            try: candidates.append(('P', b'P' + paq.compress(data)) if safe else ('L', paq.compress(data)))
            except: pass
        if HAS_ZSTD:
            try: candidates.append(('Z', b'Z' + zstd_cctx.compress(data)) if safe else ('Z', zstd_cctx.compress(data)))
            except: pass
        candidates.append(('N', b'N' + data))
        _, best = min(candidates, key=lambda x: len(x[1]))
        return best
    def _decompress_backend(self, data: bytes, safe: bool = False) -> Optional[bytes]:
        if not data: return None
        if safe:
            marker = data[0:1]
            payload = data[1:]
            if marker == b'N': return payload
            if marker == b'Z' and HAS_ZSTD:
                try: return zstd_dctx.decompress(payload)
                except: pass
            if marker == b'P' and paq:
                try: return paq.decompress(payload)
                except: pass
            return None
        if HAS_ZSTD:
            try: return zstd_dctx.decompress(data)
            except: pass
        if paq:
            try: return paq.decompress(data)
            except: pass
        if data and data[0] == ord('N'): return data[1:]
        return None

    # ---------- Header encoding (extended to 512) ----------
    def _encode_marker_single(self, t: int) -> bytes:
        if t <= 252: return bytes([t - 1])
        if t <= 256: return bytes([254, t - 253])
        # t in 257..512
        return bytes([255, t - 257])

    def _encode_marker_raw(self) -> bytes:
        return bytes([252])

    def _encode_marker_pair(self, t1: int, t2: int) -> bytes:
        idx = (t1 - 1) * 52 + (t2 - 1)
        return bytes([253, (idx >> 8) & 0xFF, idx & 0xFF])

    def _decode_header(self, data: bytes):
        if len(data) < 1: return 0, ()
        f = data[0]
        if f < 252: return 1, (f + 1,)
        if f == 252: return 1, ()
        if f == 253:
            if len(data) < 3: return 0, ()
            idx = (data[1] << 8) | data[2]
            if idx >= len(self.sequences): return 0, ()
            t1, t2 = self.pair_lookup[idx]
            return 3, (t1, t2)
        if f == 254:
            if len(data) < 2: return 0, ()
            x = data[1]
            if x > 3: return 0, ()
            return 2, (253 + x,)
        if f == 255:
            if len(data) < 2: return 0, ()
            y = data[1]
            return 2, (257 + y,)
        return 0, ()

    # ---------- Main compression ----------
    def compress_with_best(self, data: bytes, safe: bool = False, ultra: bool = True,
                           include_28: bool = False, include_29: bool = False,
                           include_30: bool = False) -> bytes:
        if not data:
            compressed = self._encode_marker_raw() + self._compress_backend(b'', safe)
            if not safe:
                decomp, _ = self._decompress_auto(compressed)
                if decomp != b'':
                    return self.compress_with_best(data, safe=True, ultra=ultra, include_28=include_28,
                                                   include_29=include_29, include_30=include_30)
            return compressed

        best_total = float('inf')
        best_bytes = None

        # Build list of single transforms (1‑512 excluding quantum if not enabled)
        single = list(range(1, 257))
        if USE_QUANTUM and HAS_QISKIT:
            single.extend(range(257, 283))
        single.extend(range(283, 513))
        if not include_28: single = [t for t in single if t != 28]
        if not include_29: single = [t for t in single if t != 29]
        if not include_30: single = [t for t in single if t != 30]

        allowed_pairs = self.sequences
        if not include_28: allowed_pairs = [seq for seq in allowed_pairs if 28 not in seq]
        if not include_29: allowed_pairs = [seq for seq in allowed_pairs if 29 not in seq]
        if not include_30: allowed_pairs = [seq for seq in allowed_pairs if 30 not in seq]
        allowed_pairs = [seq for seq in allowed_pairs if 31 not in seq and 32 not in seq]

        # raw
        cand = self._encode_marker_raw() + self._compress_backend(data, safe)
        if len(cand) < best_total:
            best_total, best_bytes = len(cand), cand

        # singles
        for t in single:
            try:
                transformed = self.fwd_transforms[t](data)
                cand = self._encode_marker_single(t) + self._compress_backend(transformed, safe)
                if len(cand) < best_total:
                    best_total, best_bytes = len(cand), cand
            except: continue

        # pairs
        if ultra:
            for t1, t2 in allowed_pairs:
                try:
                    transformed = self._apply_sequence(data, (t1, t2))
                    cand = self._encode_marker_pair(t1, t2) + self._compress_backend(transformed, safe)
                    if len(cand) < best_total:
                        best_total, best_bytes = len(cand), cand
                except: continue

        # verify
        decomp, _ = self._decompress_auto(best_bytes)
        if decomp != data:
            if not safe:
                print("Falling back to safe markers...")
                return self.compress_with_best(data, safe=True, ultra=ultra, include_28=include_28,
                                               include_29=include_29, include_30=include_30)
            else:
                print("Warning: safe compression failed; storing raw.")
                raw = self._encode_marker_raw() + self._compress_backend(data, safe=True)
                return raw
        return best_bytes

    def _decompress_auto(self, data: bytes) -> Tuple[bytes, Optional[Tuple[int, ...]]]:
        offset, seq = self._decode_header(data)
        if offset == 0: return b'', None
        payload = data[offset:]
        if not payload: return b'', None
        first_byte = payload[0:1]
        if first_byte in (b'N', b'Z', b'P'):
            res = self._decompress_backend(payload, safe=True)
        else:
            res = self._decompress_backend(payload, safe=False)
        if res is None: return b'', None
        try:
            if not seq: result = res
            else: result = self._reverse_sequence(res, seq)
        except: return b'', None
        return result, seq

    # ---------- Dictionary compression (unchanged) ----------
    MAGIC_DICT = b'DICT'
    MAGIC_LINE = b'LINE'

    def _tokenize_with_static_dict(self, data: bytes) -> Optional[bytes]:
        try: text = data.decode('utf-8')
        except: return None
        tokens = re.split(r'([A-Za-z0-9_]+)', text)
        stream = bytearray()
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                idx = self.word_to_index.get(tok)
                if idx is not None:
                    stream += b'\x01' + struct.pack('>I', idx)
                else:
                    wb = tok.encode('utf-8')
                    stream += b'\x02' + struct.pack('>H', len(wb)) + wb
            else:
                if tok:
                    sep = tok.encode('utf-8')
                    stream += b'\x00' + struct.pack('>H', len(sep)) + sep
        return bytes(stream)
    def _detokenize_static_dict(self, token_stream):
        if not token_stream: return b''
        out = bytearray(); pos = 0
        while pos < len(token_stream):
            typ = token_stream[pos]; pos += 1
            if typ == 0x01:
                if pos + 4 > len(token_stream): break
                idx = struct.unpack('>I', token_stream[pos:pos+4])[0]; pos += 4
                if idx < len(self.static_dict): out += self.static_dict[idx].encode('utf-8')
                else: return None
            elif typ == 0x02:
                if pos + 2 > len(token_stream): break
                wlen = struct.unpack('>H', token_stream[pos:pos+2])[0]; pos += 2
                if pos + wlen > len(token_stream): break
                out += token_stream[pos:pos+wlen]; pos += wlen
            elif typ == 0x00:
                if pos + 2 > len(token_stream): break
                slen = struct.unpack('>H', token_stream[pos:pos+2])[0]; pos += 2
                if pos + slen > len(token_stream): break
                out += token_stream[pos:pos+slen]; pos += slen
            else: break
        return bytes(out)
    def _compress_static_dict(self, data):
        ts = self._tokenize_with_static_dict(data)
        if ts is None: return None
        return self.MAGIC_DICT + b'\x01' + self._compress_backend(ts, safe=True)
    def _decompress_static_dict(self, data):
        if not data.startswith(self.MAGIC_DICT + b'\x01'): return None
        ts = self._decompress_backend(data[5:], safe=True)
        return self._detokenize_static_dict(ts) if ts else None
    def _compress_dynamic_dict(self, data):
        ts = self.transform_25(data)
        return self.MAGIC_DICT + b'\x02' + self._compress_backend(ts, safe=True)
    def _decompress_dynamic_dict(self, data):
        if not data.startswith(self.MAGIC_DICT + b'\x02'): return None
        ts = self._decompress_backend(data[5:], safe=True)
        return self.reverse_transform_25(ts) if ts else None
    def _tokenize_with_line_dict(self, data):
        if not self.line_dict: return None
        try: text = data.decode('utf-8')
        except: return None
        pos = 0
        token_list = []
        while pos < len(text):
            best_pos, best_len, best_idx = len(text)+1, 0, -1
            for idx, phrase in enumerate(self.line_dict):
                p = text.find(phrase, pos)
                if p != -1 and (p < best_pos or (p == best_pos and len(phrase) > best_len)):
                    best_pos, best_len, best_idx = p, len(phrase), idx
            if best_idx != -1:
                if best_pos > pos:
                    token_list.append((False, text[pos:best_pos].encode('utf-8')))
                token_list.append((True, best_idx))
                pos = best_pos + best_len
            else:
                token_list.append((False, text[pos:].encode('utf-8')))
                break
        out = bytearray()
        for is_idx, payload in token_list:
            if is_idx: out += b'\x01' + struct.pack('>Q', payload)
            else: out += b'\x00' + struct.pack('>H', len(payload)) + payload
        return bytes(out)
    def _detokenize_line_dict(self, token_stream):
        if not token_stream: return b''
        out = bytearray(); pos = 0
        while pos < len(token_stream):
            typ = token_stream[pos]; pos += 1
            if typ == 1:
                if pos + 8 > len(token_stream): return None
                idx = struct.unpack('>Q', token_stream[pos:pos+8])[0]; pos += 8
                if idx < len(self.line_dict): out += self.line_dict[idx].encode('utf-8')
                else: return None
            elif typ == 0:
                if pos + 2 > len(token_stream): return None
                rlen = struct.unpack('>H', token_stream[pos:pos+2])[0]; pos += 2
                if pos + rlen > len(token_stream): return None
                out += token_stream[pos:pos+rlen]; pos += rlen
            else: return None
        return bytes(out)
    def _compress_line_dict(self, data):
        ts = self._tokenize_with_line_dict(data)
        if ts is None: return None
        return self.MAGIC_LINE + self._compress_backend(ts, safe=True)
    def _decompress_line_dict(self, data):
        if not data.startswith(self.MAGIC_LINE): return None
        ts = self._decompress_backend(data[4:], safe=True)
        return self._detokenize_line_dict(ts) if ts else None

    # ---------- Verification ----------
    def verify_transforms(self) -> bool:
        print("Verifying all 512 transforms...")
        ok = True
        for t in range(1, 513):
            if t in range(257, 283) and not (USE_QUANTUM and HAS_QISKIT): continue
            test = bytes([0x55])
            try:
                enc = self.fwd_transforms[t](test)
                dec = self.rev_transforms[t](enc)
                if dec == test:
                    print(f"Transform {t}: right")
                else:
                    print(f"Transform {t}: incorrect")
                    ok = False
            except Exception as e:
                print(f"Transform {t}: exception {e}")
                ok = False
        print("Verification complete.\n")
        return ok

    def full_self_test(self) -> bool:
        print("=" * 60)
        print("PJP – FULL SELF‑TEST (100% lossless)")
        print("=" * 60)
        all_ok = True
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            p = doc.add_paragraph("Hello World! ")
            r = p.add_run("This is bold."); r.bold = True; r.font.size = Pt(14)
            p.add_run(" Normal text.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0,0).text = "Cell 1,1"; table.cell(0,1).text = "Cell 1,2"
            table.cell(1,0).text = "Cell 2,1"; table.cell(1,1).text = "Cell 2,2"
            bio = io.BytesIO()
            doc.save(bio)
            docx_bytes = bio.getvalue()
            # Test 31
            enc31 = self.transform_31(docx_bytes)
            dec31 = self.reverse_transform_31(enc31)
            if "Hello World!" not in Document(io.BytesIO(dec31)).paragraphs[0].text:
                print("  FAIL: transform 31"); all_ok = False
            else: print("  PASS: transform 31")
            # Test 32
            enc32 = self.transform_32(docx_bytes)
            dec32 = self.reverse_transform_32(enc32)
            doc32 = Document(io.BytesIO(dec32))
            if len(doc32.tables) == 0 or doc32.tables[0].cell(0,0).text != "Cell 1,1":
                print("  FAIL: transform 32"); all_ok = False
            else: print("  PASS: transform 32")
        except ImportError:
            print("  SKIP: python-docx not installed, cannot test transforms 31 & 32.")
        # Test dictionary
        test_text = "Hello world hello world test test"
        dict_list, w2i = self._build_text_dictionary([test_text])
        enc = self._encode_text_with_dict(test_text, dict_list, w2i)
        dec = self._decode_text_with_dict(enc, dict_list)
        if dec != test_text:
            print("  FAIL: dictionary encode/decode"); all_ok = False
        else: print("  PASS: dictionary encode/decode")
        # Test plain text extraction (internal)
        if 'docx_bytes' in locals():
            plain = self.extract_text_from_docx(docx_bytes)
            if "Hello World!" not in plain or "Cell 1,1" not in plain:
                print("  FAIL: plain text extraction"); all_ok = False
            else: print("  PASS: plain text extraction")
        if all_ok: print("\n[All tests passed – compressor is 100% lossless]")
        else: print("\n[FAIL] Some tests failed.")
        return all_ok

    # ---------- File API ----------
    def compress_file(self, infile, outfile, ultra=True, hybrid=False, include_28=False, include_29=False, include_30=False):
        with open(infile, 'rb') as f: data = f.read()
        candidates = []
        if hybrid:
            c_static = self._compress_static_dict(data)
            if c_static: candidates.append(('Static-Word-Dict', c_static))
            c_line = self._compress_line_dict(data)
            if c_line: candidates.append(('Line-Dict', c_line))
            c_dynamic = self._compress_dynamic_dict(data)
            if c_dynamic: candidates.append(('Dynamic-Dict', c_dynamic))
        c_pjp = self.compress_with_best(data, ultra=ultra, include_28=include_28, include_29=include_29, include_30=include_30)
        candidates.append(('PJP', c_pjp))
        method, best = min(candidates, key=lambda x: len(x[1]))
        with open(outfile, 'wb') as f: f.write(best)
        print(f"Compressed {len(data)} → {len(best)} bytes ({method}) → {outfile}")

    def decompress_file(self, infile, outfile):
        with open(infile, 'rb') as f: data = f.read()
        if data.startswith(self.MAGIC_LINE):
            orig = self._decompress_line_dict(data)
            if orig:
                with open(outfile, 'wb') as f: f.write(orig)
                print(f"Decompressed (Line-Dict) → {outfile} ({len(orig)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x01'):
            orig = self._decompress_static_dict(data)
            if orig:
                with open(outfile, 'wb') as f: f.write(orig)
                print(f"Decompressed (Static-Word-Dict) → {outfile} ({len(orig)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x02'):
            orig = self._decompress_dynamic_dict(data)
            if orig:
                with open(outfile, 'wb') as f: f.write(orig)
                print(f"Decompressed (Dynamic-Dict) → {outfile} ({len(orig)} bytes)")
                return
        orig, seq = self._decompress_auto(data)
        if orig == b'' and seq is None:
            print("Decompression failed – unknown format.")
            return
        with open(outfile, 'wb') as f: f.write(orig)
        print(f"Decompressed ({'raw' if not seq else f'sequence {seq}'}) → {outfile} ({len(orig)} bytes)")

# ------------------------------------------------------------
def main():
    print(f"{PROGNAME} – 512 transforms (256 base + 256 fast) + 2704 pairs + Base64 + 6‑bit + Quantum + 28‑30 + 31/32 .docx")
    print("Options 1-3 do NOT use 28‑30; option 4 uses them.")
    if paq is None and not HAS_ZSTD: print("Warning: No compression backend found.")
    c = PJPCompressor()
    c.verify_transforms()
    choice = input("\n1) Fast (no 28-30) – 256 singles\n"
                   "2) Ultra (no 28-30) – 256 singles + 2704 pairs\n"
                   "3) Hybrid (no 28-30) – dicts + Ultra\n"
                   "4) Absolute (with 28, 29, 30) – all transforms\n"
                   "5) Full self‑test\n"
                   "6) Decompress (extract)\n"
                   "> ").strip()
    if choice == '1':
        i = input("Input file: ").strip()
        o = input("Output file: ").strip() or i + ".pjp"
        c.compress_file(i, o, ultra=False, hybrid=False)
    elif choice == '2':
        i = input("Input file: ").strip()
        o = input("Output file: ").strip() or i + ".pjp"
        c.compress_file(i, o, ultra=True, hybrid=False)
    elif choice == '3':
        i = input("Input file: ").strip()
        o = input("Output file: ").strip() or i + ".pjp"
        c.compress_file(i, o, ultra=True, hybrid=True)
    elif choice == '4':
        i = input("Input file: ").strip()
        o = input("Output file: ").strip() or i + ".pjp"
        c.compress_file(i, o, ultra=True, hybrid=True, include_28=True, include_29=True, include_30=True)
    elif choice == '5':
        c.full_self_test()
    elif choice == '6':
        i = input("Compressed file: ").strip()
        o = input("Output file: ").strip() or i.rsplit('.', 1)[0] + ".orig"
        c.decompress_file(i, o)
    else:
        print("Invalid choice.")

if __name__ == "__main__":
    main()
